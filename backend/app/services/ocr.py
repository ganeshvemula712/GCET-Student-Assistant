import logging
import os
import re
import time
from io import BytesIO

from backend.app.core.config import settings
from backend.app.services.gemini import client as gemini_client
from google.genai import types

logger = logging.getLogger("uvicorn")


class GeminiQuotaExhaustedError(Exception):
    """Raised when Gemini daily quota / free tier limit is exhausted."""
    pass


class GeminiRateLimitError(Exception):
    """Raised when temporary 429 RPM rate limit occurs."""
    pass


def is_daily_quota_error(err_str: str) -> bool:
    err_lower = err_str.lower()
    return any(
        indicator in err_lower
        for indicator in [
            "generate_content_free_tier_requests",
            "quota exceeded for metric",
            "daily_quota",
            "resourcehasbeenexhausted",
            "free_tier_requests",
        ]
    )


def call_gemini_ocr_batch_with_retry(
    image_parts_with_page_nums: list[tuple[int, types.Part]],
    max_retries: int = 2,
) -> list[tuple[int, str]]:
    """
    Batches multiple image parts into a single Gemini Vision API call.
    Returns a list of (page_number, extracted_text).
    """
    if not image_parts_with_page_nums:
        return []

    page_markers = ", ".join([f"Page {p_num}" for p_num, _ in image_parts_with_page_nums])
    ocr_prompt = (
        f"You are an expert OCR and document analysis engine for GCET College.\n"
        f"Below are {len(image_parts_with_page_nums)} scanned document page image(s) corresponding to [{page_markers}].\n"
        f"Extract all text, notices, examination guidelines, and timetable schedules from each image.\n"
        f"For timetables, format them as clean Markdown tables with Day, Time, Subject, Room, and Faculty columns.\n"
        f"Preserve exact course names, exam dates, and timing.\n"
        f"CRITICAL: Separate each page's extracted text clearly using the exact header format:\n"
        f"--- PAGE <page_number> ---\n"
        f"followed by the extracted text for that specific page."
    )

    contents = []
    for p_num, part in image_parts_with_page_nums:
        contents.append(part)
    contents.append(ocr_prompt)

    attempt = 0
    response_text = ""

    while attempt <= max_retries:
        try:
            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
            )
            response_text = response.text or ""
            break
        except Exception as err:
            err_str = str(err)
            logger.warning(f"[OCR] Gemini API call attempt {attempt + 1} failed: {err_str}")
            if is_daily_quota_error(err_str):
                logger.error(f"[OCR CRITICAL] Gemini daily quota exhausted: {err_str}")
                raise GeminiQuotaExhaustedError(
                    "Document uploaded successfully, but OCR indexing could not be completed because the Gemini OCR quota is currently exhausted. Please retry after the quota resets."
                ) from err
            elif "429" in err_str or "resource_exhausted" in err_str.lower():
                attempt += 1
                if attempt <= max_retries:
                    backoff = 2 * attempt
                    logger.info(f"[OCR] Temporary rate limit encountered. Retrying in {backoff}s...")
                    time.sleep(backoff)
                else:
                    raise GeminiRateLimitError(f"Temporary Gemini rate limit exceeded after {max_retries} retries.") from err
            else:
                raise err

    if not response_text.strip():
        return []

    # Parse response text back into individual pages
    results = []
    page_blocks = re.split(r"--- PAGE\s+(\d+)\s+---", response_text, flags=re.IGNORECASE)

    if len(page_blocks) > 1:
        for i in range(1, len(page_blocks), 2):
            try:
                p_num = int(page_blocks[i].strip())
                p_text = page_blocks[i + 1].strip()
                if p_text:
                    results.append((p_num, p_text))
            except (ValueError, IndexError):
                continue
    else:
        first_page_num = image_parts_with_page_nums[0][0]
        results.append((first_page_num, response_text.strip()))

    return results


import gc

def extract_document_pages(content: bytes, filename: str) -> tuple[int, list[tuple[int, str]]]:
    """
    Extracts text from PDF, DOCX, or Image files.
    PDF logic:
    1. Checks each page for native text (`page.get_text()`).
    2. If native text length >= 20, extracts natively with ZERO Gemini OCR calls.
    3. Scanned/image pages are rendered and batched in groups of OCR_BATCH_SIZE
       streamed page-by-page to prevent RAM spikes on 512MB RAM instances.
    """
    filename_lower = filename.lower()
    extracted_pages = []
    page_count = 1

    if filename_lower.endswith(".pdf"):
        import pymupdf
        pdf_doc = pymupdf.open(stream=BytesIO(content), filetype="pdf")
        page_count = len(pdf_doc)

        scanned_page_indices = []

        for page_index in range(page_count):
            page = pdf_doc[page_index]
            text = page.get_text()
            if text and len(text.strip()) >= 20:
                extracted_pages.append((page_index + 1, text.strip()))
            else:
                scanned_page_indices.append(page_index)

        batch_size = getattr(settings, "OCR_BATCH_SIZE", 4)
        for i in range(0, len(scanned_page_indices), batch_size):
            batch_indices = scanned_page_indices[i : i + batch_size]
            batch_parts = []
            for p_idx in batch_indices:
                page = pdf_doc[p_idx]
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("jpeg")
                image_part = types.Part.from_bytes(data=img_bytes, mime_type="image/jpeg")
                batch_parts.append((p_idx + 1, image_part))
                del pix

            ocr_results = call_gemini_ocr_batch_with_retry(batch_parts)
            extracted_pages.extend(ocr_results)
            del batch_parts
            gc.collect()

        pdf_doc.close()

    elif filename_lower.endswith((".docx", ".doc")):
        import docx
        doc_obj = docx.Document(BytesIO(content))
        lines = []
        for p in doc_obj.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        for table in doc_obj.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                lines.append("\n" + "\n".join(table_lines) + "\n")
        full_text = "\n\n".join(lines)
        if full_text.strip():
            extracted_pages.append((1, full_text))
        page_count = 1

    elif filename_lower.endswith((".jpg", ".jpeg", ".png")):
        mime_type = "image/jpeg" if filename_lower.endswith((".jpg", ".jpeg")) else "image/png"
        image_part = types.Part.from_bytes(data=content, mime_type=mime_type)
        ocr_results = call_gemini_ocr_batch_with_retry([(1, image_part)])
        extracted_pages.extend(ocr_results)
        del image_part
        gc.collect()
        page_count = 1

    extracted_pages.sort(key=lambda x: x[0])
    return page_count, extracted_pages

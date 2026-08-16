import asyncio
import logging
from io import BytesIO

from backend.app.core.database import SessionLocal
from backend.app.models.document import Document
from backend.app.services.embeddings import generate_embeddings, _cached_query_embedding
from backend.app.services.gemini import client as gemini_client
from backend.app.services.storage import download_file_from_storage, get_storage_key
from backend.app.services.text_processing import chunk_text
from backend.app.services.vector_store import get_collection, store_chunks
from google.genai import types

logger = logging.getLogger("uvicorn")


def _reindex_document_from_bytes(doc: Document, content: bytes) -> int:
    filename_lower = doc.filename.lower()
    page_count = 1
    extracted_pages = []

    if filename_lower.endswith(".pdf"):
        import pymupdf
        pdf_document = pymupdf.open(stream=BytesIO(content), filetype="pdf")
        page_count = len(pdf_document)
        for page_index in range(page_count):
            page = pdf_document[page_index]
            text = page.get_text()
            if text.strip():
                extracted_pages.append((page_index + 1, text))
            else:
                try:
                    pix = page.get_pixmap(dpi=150)
                    img_bytes = pix.tobytes("jpeg")
                    image_part = types.Part.from_bytes(data=img_bytes, mime_type="image/jpeg")
                    ocr_prompt = (
                        "You are an expert OCR and document analysis engine for GCET College. "
                        "Extract all text, notices, examination guidelines, and timetable schedules from this document image. "
                        "For timetables, format them as clean, structured Markdown tables with Day, Time, Subject, Room, and Faculty columns. "
                        "Preserve exact course names, exam dates, and timing."
                    )
                    response = gemini_client.models.generate_content(
                        model="gemini-2.5-flash", contents=[image_part, ocr_prompt]
                    )
                    ocr_text = response.text or ""
                    if ocr_text.strip():
                        extracted_pages.append((page_index + 1, ocr_text))
                except Exception as ocr_err:
                    logger.error(f"[SELF-HEAL] OCR failed for page {page_index + 1}: {ocr_err}")

    elif filename_lower.endswith((".docx", ".doc")):
        import docx
        doc_obj = docx.Document(BytesIO(content))
        lines = []
        for p in doc_obj.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        for table in doc_obj.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                lines.append("\n" + "\n".join(table_lines) + "\n")
        full_text = "\n\n".join(lines)
        if full_text.strip():
            extracted_pages.append((1, full_text))

    elif filename_lower.endswith((".jpg", ".jpeg", ".png")):
        mime_type = "image/jpeg" if filename_lower.endswith((".jpg", ".jpeg")) else "image/png"
        image_part = types.Part.from_bytes(data=content, mime_type=mime_type)
        ocr_prompt = (
            "You are an expert OCR and document analysis engine for GCET College. "
            "Extract all text, notices, examination guidelines, and timetable schedules from this document image. "
            "For timetables, format them as clean, structured Markdown tables with Day, Time, Subject, Room, and Faculty columns. "
            "Preserve exact course names, exam dates, and timing."
        )
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash", contents=[image_part, ocr_prompt]
        )
        extracted_text = response.text or ""
        if extracted_text.strip():
            extracted_pages.append((1, extracted_text))

    all_chunks = []
    for p_num, p_text in extracted_pages:
        page_chunks = chunk_text(
            text=p_text,
            filename=doc.filename,
            page_number=p_num,
            document_id=doc.document_id,
            chunk_size=1000,
            chunk_overlap=200,
            category=doc.category or "General Academic",
            tags=doc.tags or "",
        )
        for chunk in page_chunks:
            chunk["metadata"]["version"] = doc.version or 1
            chunk["metadata"]["is_active"] = True
            chunk["metadata"]["category"] = doc.category or "General Academic"
            chunk["metadata"]["tags"] = doc.tags or ""
            all_chunks.append(chunk)

    if not all_chunks:
        return 0

    texts = [chunk["text"] for chunk in all_chunks]
    embeddings = generate_embeddings(texts)
    store_chunks(chunks=all_chunks, embeddings=embeddings)
    return len(all_chunks)


def sync_chromadb_with_postgres() -> None:
    """
    Self-healing background worker:
    Scans active PostgreSQL document records against ChromaDB vector storage.
    If vectors are missing (e.g. after ephemeral disk wipe), downloads source file
    from Cloudflare R2 storage and re-indexes into ChromaDB automatically.
    """
    logger.info("[SELF-HEAL] Starting ChromaDB vector persistence check...")
    db = SessionLocal()
    try:
        active_docs = db.query(Document).filter(Document.is_active == True).all()
        logger.info(f"[SELF-HEAL] Found {len(active_docs)} active document records in PostgreSQL.")

        coll = get_collection()

        for doc in active_docs:
            try:
                res = coll.get(where={"document_id": doc.document_id})
                chunk_count = len(res.get("ids", []))

                if chunk_count > 0:
                    logger.info(f"[SELF-HEAL] Document '{doc.filename}' ({doc.document_id[:8]}) verified: {chunk_count} vectors in ChromaDB.")
                    if doc.status != "processed":
                        doc.status = "processed"
                else:
                    logger.warning(f"[SELF-HEAL] Vectors missing for active document '{doc.filename}' ({doc.document_id[:8]}). Healing from storage...")
                    obj_key = get_storage_key(doc.document_id, doc.filename)
                    content = download_file_from_storage(obj_key)

                    if content:
                        new_count = _reindex_document_from_bytes(doc, content)
                        if new_count > 0:
                            doc.status = "processed"
                            doc.chunk_count = new_count
                            logger.info(f"[SELF-HEAL] Successfully restored '{doc.filename}': {new_count} chunks indexed into ChromaDB.")
                        else:
                            doc.status = "indexing_required"
                            logger.error(f"[SELF-HEAL] Extraction empty for '{doc.filename}'. Status set to indexing_required.")
                    else:
                        doc.status = "indexing_required"
                        logger.warning(f"[SELF-HEAL] Source file for '{doc.filename}' not found in cloud storage. Marked indexing_required.")
            except Exception as doc_err:
                logger.error(f"[SELF-HEAL] Error checking/healing document '{doc.filename}': {doc_err}")

        db.commit()
        _cached_query_embedding.cache_clear()
        logger.info("[SELF-HEAL] Vector persistence check completed successfully.")
    except Exception as err:
        db.rollback()
        logger.error(f"[SELF-HEAL] Critical error during ChromaDB synchronization: {err}")
    finally:
        db.close()


async def run_async_self_healing():
    """Run self-healing task in background without blocking server startup."""
    await asyncio.sleep(1)
    await asyncio.to_thread(sync_chromadb_with_postgres)

from hashlib import sha256
from io import BytesIO

from fastapi import HTTPException, UploadFile
from sqlalchemy import or_
from sqlalchemy.orm import Session

from backend.app.models.document import Document, DEFAULT_CATEGORY, VALID_CATEGORIES
from backend.app.services.embeddings import generate_embeddings
from backend.app.services.text_processing import chunk_text
from backend.app.services.gemini import client as gemini_client
from backend.app.services.vector_store import (
    delete_document_chunks,
    mark_document_chunks_inactive,
    store_chunks,
)
from google.genai import types


def normalize_tags(tags_input: str | None) -> str:
    if not tags_input:
        return ""
    raw_tags = [t.strip() for t in tags_input.split(",")]
    seen = set()
    clean_tags = []
    for t in raw_tags:
        if t and t.lower() not in seen:
            seen.add(t.lower())
            clean_tags.append(t)
    return ", ".join(clean_tags)


async def process_document(
    file: UploadFile,
    db: Session,
    supersedes_id: str | None = None,
    category: str | None = None,
    tags: str | None = None,
) -> dict:
    filename_lower = file.filename.lower()
    allowed_extensions = (".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png")

    if not any(filename_lower.endswith(ext) for ext in allowed_extensions):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Supported formats: PDF, DOCX, DOC, JPG, JPEG, PNG.",
        )

    # Category Validation & Tag Normalization
    category_clean = (category or "").strip()
    if category_clean and category_clean not in VALID_CATEGORIES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid category '{category_clean}'. Supported categories: {', '.join(VALID_CATEGORIES)}",
        )
    final_category = category_clean if category_clean else DEFAULT_CATEGORY
    final_tags = normalize_tags(tags)

    # Read uploaded file
    content = await file.read()

    if not content:
        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty",
        )

    # Generate unique document ID
    document_id = sha256(content).hexdigest()

    # Check whether document already exists
    existing_document = (
        db.query(Document)
        .filter(Document.document_id == document_id)
        .first()
    )

    if existing_document:
        raise HTTPException(
            status_code=409,
            detail="This document has already been uploaded",
        )

    page_count = 1
    extracted_pages = []

    # 1. Process PDF
    if filename_lower.endswith(".pdf"):
        try:
            import pymupdf
            pdf_document = pymupdf.open(
                stream=BytesIO(content),
                filetype="pdf",
            )
            page_count = len(pdf_document)
            for page_index in range(page_count):
                page = pdf_document[page_index]
                text = page.get_text()
                if text.strip():
                    extracted_pages.append((page_index + 1, text))
                else:
                    # Fallback to Gemini Vision OCR for scanned/image-only PDF pages
                    try:
                        pix = page.get_pixmap(dpi=150)
                        img_bytes = pix.tobytes("jpeg")
                        image_part = types.Part.from_bytes(data=img_bytes, mime_type="image/jpeg")

                        ocr_prompt = (
                            "You are an expert OCR and document analysis engine for GCET College. "
                            "Extract all text, notices, examination guidelines, and timetable schedules from this document image. "
                            "For timetables, format them as clean, structured Markdown tables with Day, Time, Subject, Room, and Faculty columns. "
                            "Preserve exact course names, exam dates, and timing."
                        )

                        response = gemini_client.models.generate_content(
                            model="gemini-2.5-flash",
                            contents=[image_part, ocr_prompt],
                        )
                        ocr_text = response.text or ""
                        if ocr_text.strip():
                            extracted_pages.append((page_index + 1, ocr_text))
                    except Exception as ocr_err:
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to process OCR for scanned PDF page {page_index + 1}: {str(ocr_err)}",
                        )
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(
                status_code=400,
                detail="Failed to parse PDF document",
            )

    # 2. Process DOCX / DOC
    elif filename_lower.endswith((".docx", ".doc")):
        try:
            import docx
            doc_obj = docx.Document(BytesIO(content))
            lines = []
            for p in doc_obj.paragraphs:
                if p.text.strip():
                    lines.append(p.text)
            for table in doc_obj.tables:
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_lines.append(" | ".join(cells))
                if table_lines:
                    lines.append("\n" + "\n".join(table_lines) + "\n")
            full_text = "\n\n".join(lines)
            if full_text.strip():
                extracted_pages.append((1, full_text))
            page_count = 1
        except Exception:
            raise HTTPException(
                status_code=400,
                detail="Failed to parse Word (.docx/.doc) document",
            )

    # 3. Process Images (JPG, JPEG, PNG) via Gemini Multimodal OCR Vision
    elif filename_lower.endswith((".jpg", ".jpeg", ".png")):
        try:
            from PIL import Image
            mime_type = "image/jpeg" if filename_lower.endswith((".jpg", ".jpeg")) else "image/png"
            image_part = types.Part.from_bytes(data=content, mime_type=mime_type)

            ocr_prompt = (
                "You are an expert OCR and document analysis engine for GCET College. "
                "Extract all text, notices, examination guidelines, and timetable schedules from this document image. "
                "For timetables, format them as clean, structured Markdown tables with Day, Time, Subject, Room, and Faculty columns. "
                "Preserve exact course names, exam dates, and timing."
            )

            response = gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[image_part, ocr_prompt],
            )
            extracted_text = response.text or ""
            if extracted_text.strip():
                extracted_pages.append((1, extracted_text))
            page_count = 1
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process image OCR via Gemini Vision: {str(e)}",
            )

    # Handle Versioning
    new_version = 1
    if supersedes_id:
        old_doc = (
            db.query(Document)
            .filter(Document.document_id == supersedes_id)
            .first()
        )
        if old_doc:
            new_version = (old_doc.version or 1) + 1
            old_doc.is_active = False
            mark_document_chunks_inactive(old_doc.document_id)

    all_chunks = []
    for p_num, p_text in extracted_pages:
        page_chunks = chunk_text(
            text=p_text,
            filename=file.filename,
            page_number=p_num,
            document_id=document_id,
            chunk_size=1000,
            chunk_overlap=200,
            category=final_category,
            tags=final_tags,
        )
        for chunk in page_chunks:
            chunk["metadata"]["version"] = new_version
            chunk["metadata"]["is_active"] = True
            chunk["metadata"]["category"] = final_category
            chunk["metadata"]["tags"] = final_tags
            all_chunks.append(chunk)

    if not all_chunks:
        raise HTTPException(
            status_code=400,
            detail="No readable text or content found in uploaded document",
        )

    # Generate Embeddings
    try:
        texts = [chunk["text"] for chunk in all_chunks]
        embeddings = generate_embeddings(texts)
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Embedding generation failed",
        )

    # Store Chunks in ChromaDB
    try:
        store_chunks(
            chunks=all_chunks,
            embeddings=embeddings,
        )
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Failed to store vectors in ChromaDB",
        )

    # Save to PostgreSQL
    new_document = Document(
        document_id=document_id,
        filename=file.filename,
        page_count=page_count,
        chunk_count=len(all_chunks),
        status="processed",
        version=new_version,
        is_active=True,
        supersedes_id=supersedes_id,
        category=final_category,
        tags=final_tags,
    )

    try:
        db.add(new_document)
        db.commit()
        db.refresh(new_document)

    except Exception:
        db.rollback()
        delete_document_chunks(document_id)
        raise HTTPException(
            status_code=500,
            detail="Failed to save document record",
        )

    return {
        "id": new_document.id,
        "document_id": new_document.document_id,
        "filename": new_document.filename,
        "page_count": new_document.page_count,
        "chunk_count": new_document.chunk_count,
        "status": new_document.status,
        "version": new_document.version,
        "is_active": new_document.is_active,
        "supersedes_id": new_document.supersedes_id,
        "category": new_document.category or DEFAULT_CATEGORY,
        "tags": new_document.tags or "",
        "uploaded_at": new_document.uploaded_at,
        "message": f"Document '{file.filename}' (v{new_version}) categorized as '{final_category}' processed and indexed successfully into ChromaDB!",
    }


def get_documents(
    db: Session,
    category: str | None = None,
) -> list[Document]:
    query = db.query(Document)
    if category:
        cat_clean = category.strip()
        if cat_clean.lower() == DEFAULT_CATEGORY.lower():
            query = query.filter(
                or_(
                    Document.category.ilike(cat_clean),
                    Document.category.is_(None),
                )
            )
        else:
            query = query.filter(Document.category.ilike(cat_clean))
    return query.order_by(Document.uploaded_at.desc()).all()


def remove_document(
    document_id: str,
    db: Session,
) -> dict:
    document = (
        db.query(Document)
        .filter(Document.document_id == document_id)
        .first()
    )

    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found",
        )

    filename = document.filename
    delete_document_chunks(document_id)

    db.delete(document)
    db.commit()

    return {
        "message": "Document deleted successfully",
        "document_id": document_id,
        "filename": filename,
    }